
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Criar documento
doc = Document()

# Configurar margens (3cm esquerda/superior, 2cm direita/inferior)
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# Configurar estilo Normal
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.font.color.rgb = RGBColor(0, 0, 0)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.first_line_indent = Cm(1.25)

# Funções auxiliares
def add_para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, indent=Cm(1.25), space_after=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.first_line_indent = indent
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rPr.insert(0, rFonts)
    return p

def add_title(doc, text):
    return add_para(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, indent=Cm(0), space_after=Pt(6))

def add_center(doc, text, bold=False):
    return add_para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=bold, indent=Cm(0))

def page_break(doc):
    doc.add_page_break()

print("Documento e funções criados.")