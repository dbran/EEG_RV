from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# Criar documento
doc = Document()

# Configurar margens (3cm esquerda/superior, 2cm direita/inferior)
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# Configurar estilo Normal
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.font.color.rgb = RGBColor(0, 0, 0)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.first_line_indent = Cm(1.25)

# Funções auxiliares
def add_para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, indent=Cm(1.25), space_after=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.first_line_indent = indent
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rPr.insert(0, rFonts)
    return p

def add_title(doc, text):
    return add_para(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, indent=Cm(0), space_after=Pt(6))

def add_center(doc, text, bold=False):
    return add_para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=bold, indent=Cm(0))

def page_break(doc):
    doc.add_page_break()

# ============================================
# 8 CONSIDERAÇÕES ÉTICAS
# ============================================
add_title(doc, "8 CONSIDERAÇÕES ÉTICAS")
add_para(doc, "")

add_para(doc, "O presente estudo será submetido à apreciação pelo Comitê de Ética em Pesquisa (CEP) da Universidade "
              "Federal de Ciências da Saúde de Porto Alegre, conforme os termos da Resolução nº 466, de 12 de dezembro "
              "de 2012, do Conselho Nacional de Saúde (CNS), e da Lei nº 11.794, de 8 de outubro de 2008, que regulam "
              "as pesquisas envolvendo seres humanos no Brasil.")

add_para(doc, "A fase de validação com participantes saudáveis compreenderá: (i) Termo de Consentimento Livre e "
              "Esclarecido (TCLE): documento detalhado informando os objetivos, procedimentos, riscos, benefícios e "
              "garantias de sigilo e anonimato; (ii) Riscos minimizados: os procedimentos de EEG não invasivo apresentam "
              "riscos mínimos (possível desconforto leve dos eletrodos). A exposição à VR será limitada a sessões de "
              "30 minutos para prevenção de ciberenjoo; (iii) Sigilo e anonimato: todos os dados serão anonimizados e "
              "armazenados em ambiente seguro, com acesso restrito à equipe de pesquisa; (iv) Benefícios: contribuição "
              "para o avanço do conhecimento em neuroreabilitação assistida por tecnologia; possibilidade de acesso "
              "futuro a sistemas de reabilitação inovadores; e (v) Retirada voluntária: os participantes poderão desistir "
              "a qualquer momento, sem qualquer prejuízo.")

add_para(doc, "O referido projeto se enquadra nos termos da Resolução 466/12 do CONEP e da Lei nº 11.794 de 8 de "
              "outubro de 2008, que regulam CEP e CONEP.")

page_break(doc)

# ============================================
# 9 CONSIDERAÇÕES SOBRE RISCOS
# ============================================
add_title(doc, "9 CONSIDERAÇÕES SOBRE RISCOS")
add_para(doc, "")

riscos = [
    "Indisponibilidade de equipamento de EEG – Probabilidade: Média – Impacto: Alto – Mitigação: parceria com laboratórios parceiros; uso de simuladores de sinal",
    "Latência superior à meta de 150 ms – Probabilidade: Média – Impacto: Alto – Mitigação: otimização de pipeline; uso de hardware dedicado; paralelização de threads",
    "Baixa acurácia de classificação (menor que 75%) – Probabilidade: Média – Impacto: Alto – Mitigação: aumento do dataset de treinamento; técnicas de transfer learning; ensemble de modelos",
    "Compatibilidade limitada entre engines VR – Probabilidade: Baixa – Impacto: Médio – Mitigação: uso de APIs padronizadas (OpenXR); testes precoce em múltiplas plataformas",
    "Ciberenjoo em participantes durante VR – Probabilidade: Baixa – Impacto: Médio – Mitigação: sessões curtas (menor ou igual a 30 min); taxa de refresh maior ou igual a 90 Hz; mecanismos de conforto visual",
    "Atraso na aprovação ética – Probabilidade: Baixa – Impacto: Médio – Mitigação: submissão antecipada; diálogo contínuo com o CEP",
    "Não aprovação do CEP – Probabilidade: Baixa – Impacto: Alto – Mitigação: TCLE robusto; submissão antecipada; diálogo com CEP",
    "Inviabilidade econômica no SUS – Probabilidade: Média – Impacto: Médio – Mitigação: análise de custo-efetividade; parcerias com gestores de saúde"
]

for r in riscos:
    add_para(doc, r)

page_break(doc)

# ============================================
# REFERÊNCIAS BIBLIOGRÁFICAS
# ============================================
add_title(doc, "REFERÊNCIAS BIBLIOGRÁFICAS")
add_para(doc, "")

referencias = [
    "ANG, K. K. et al. A randomized controlled trial of EEG-based motor imagery brain-computer interface robotic rehabilitation for stroke. Clinical EEG and Neuroscience, v. 46, n. 4, p. 310-320, 2015. Disponível em: https://pubmed.ncbi.nlm.nih.gov/25055897/. Acesso em: 18 jul. 2026.",
    
    "BAHIA, L. et al. Costs of stroke in Brazil: protocol for a societal cost-of-illness study. Value in Health Regional Issues, v. 22, p. 36-43, 2020. Disponível em: https://www.sciencedirect.com/science/article/pii/S2212109920300224. Acesso em: 18 jul. 2026.",
    
    "BLANKERTZ, B. et al. Optimizing spatial filters for robust EEG single-trial analysis. IEEE Signal Processing Magazine, v. 25, n. 1, p. 41-56, 2008. Disponível em: https://ieeexplore.ieee.org/document/4408441. Acesso em: 18 jul. 2026.",
    
    "BONINGER, M. L. et al. Technology and disability: findings from a National Institute on Disability and Rehabilitation Research conference. Journal of Rehabilitation Research & Development, v. 51, n. 6, p. 855-866, 2014. Disponível em: https://pubmed.ncbi.nlm.nih.gov/25356922/. Acesso em: 18 jul. 2026.",
    
    "CARAMIA, M. D. et al. Brain-computer interface for hand motor function rehabilitation in stroke: a systematic review and meta-analysis. Frontiers in Neuroscience, v. 15, p. 642732, 2021. Disponível em: https://www.frontiersin.org/articles/10.3389/fnins.2021.642732/full. Acesso em: 18 jul. 2026.",
    
    "DIETRICH, C. Neuroplasticity and motor recovery after stroke: evidence from functional neuroimaging. Restorative Neurology and Neuroscience, v. 39, n. 3, p. 185-196, 2021. Disponível em: https://pubmed.ncbi.nlm.nih.gov/33950291/. Acesso em: 18 jul. 2026.",
    
    "FEIGIN, V. L. et al. Global, regional, and national burden of stroke and its risk factors, 1990–2019: a systematic analysis for the Global Burden of Disease Study 2019. The Lancet Neurology, v. 20, n. 10, p. 795-820, 2021. Disponível em: https://www.thelancet.com/journals/laneur/article/PIIS1474-4422(21)00252-0/fulltext. Acesso em: 18 jul. 2026.",
    
    "FUGL-MEYER, A. R. et al. The post-stroke hemiplegic patient: a method for evaluation of physical performance. Scandinavian Journal of Rehabilitation Medicine, v. 7, n. 1, p. 13-31, 1975. Disponível em: https://pubmed.ncbi.nlm.nih.gov/1135616/. Acesso em: 18 jul. 2026.",
    
    "GBD 2019 STROKE COLLABORATORS. Global, regional, and national burden of stroke and its risk factors, 1990-2019: a systematic analysis for the Global Burden of Disease Study 2019. The Lancet Neurology, v. 20, n. 10, p. 795-820, 2021. Disponível em: https://www.thelancet.com/journals/laneur/article/PIIS1474-4422(21)00252-0/fulltext. Acesso em: 18 jul. 2026.",
    
    "GORMAN, C. et al. A closed-loop AR-based BCI for real-world system control. In: IEEE Symposium Series on Computational Intelligence (SSCI). 2021. Disponível em: https://ieeexplore.ieee.org/xpl/conhome/9651419/proceeding. Acesso em: 18 jul. 2026.",
    
    "HAL. Designing functional prototypes combining BCI and AR. EuroXR Conference, 2023. Disponível em: https://hal.science/hal-03928273. Acesso em: 18 jul. 2026.",
    
    "HTC VIVE. OpenXR hand tracking SDK tutorial. VIVE Developer, 2024. Disponível em: https://developer.vive.com/resources/openxr/. Acesso em: 18 jul. 2026.",
    
    "KELLEY, R. E. et al. Functional recovery after stroke: a review of current therapeutic strategies. Current Neurology and Neuroscience Reports, v. 23, n. 5, p. 245-257, 2023. Disponível em: https://pubmed.ncbi.nlm.nih.gov/37192684/. Acesso em: 18 jul. 2026.",
    
    "KLEIM, J. A.; JONES, T. A. Principles of experience-dependent neural plasticity: implications for rehabilitation after brain damage. Journal of Speech, Language, and Hearing Research, v. 51, n. 1, p. S225-S239, 2008. Disponível em: https://pubmed.ncbi.nlm.nih.gov/18230848/. Acesso em: 18 jul. 2026.",
    
    "KWAKKEL, G. et al. Effects of augmented exercise therapy time after stroke: a meta-analysis. Stroke, v. 35, n. 11, p. 2529-2539, 2004. Disponível em: https://pubmed.ncbi.nlm.nih.gov/15472114/. Acesso em: 18 jul. 2026.",
    
    "LANG, C. E. et al. Upper extremity use in people with hemiparesis in the first six weeks after stroke. Journal of Neurologic Physical Therapy, v. 31, n. 2, p. 56-63, 2020. Disponível em: https://pubmed.ncbi.nlm.nih.gov/17534167/. Acesso em: 18 jul. 2026.",
    
    "LAWHERN, V. J. et al. EEGNet: a compact convolutional neural network for EEG-based brain-computer interfaces. Journal of Neural Engineering, v. 15, n. 5, p. 056013, 2018. Disponível em: https://iopscience.iop.org/article/10.1088/1741-2552/aace8c. Acesso em: 18 jul. 2026.",
    
    "LAWRENCE, E. S. et al. Estimates of the prevalence of acute stroke impairments and disability in a multiethnic population. Stroke, v. 32, n. 6, p. 1279-1284, 2001. Disponível em: https://pubmed.ncbi.nlm.nih.gov/11387487/. Acesso em: 18 jul. 2026.",
    
    "LECUYER, A. et al. Brain-computer interfaces, virtual reality, and videogames. IEEE Computer, v. 41, n. 10, p. 66-72, 2008. Disponível em: https://ieeexplore.ieee.org/document/4664317. Acesso em: 18 jul. 2026.",
    
    "LI, A. et al. Interactive and deep learning-powered EEG-BCI for wrist rehabilitation: a game-based prototype study. Neurobiology, v. 9, n. 3, p. 302, 2025. Disponível em: https://www.mdpi.com/journal/neurobiology. Acesso em: 18 jul. 2026.",
    
    "LO, A. C. et al. Robot-assisted therapy for long-term upper-limb impairment after stroke. New England Journal of Medicine, v. 362, n. 19, p. 1772-1783, 2010. Disponível em: https://pubmed.ncbi.nlm.nih.gov/20400552/. Acesso em: 18 jul. 2026.",
    
    "LOTTE, F. et al. Combining BCI with virtual reality: towards new applications and improved BCI. In: Towards Practical Brain-Computer Interfaces. Springer, 2012. p. 303-331. Disponível em: https://link.springer.com/chapter/10.1007/978-3-642-29746-5_14. Acesso em: 18 jul. 2026.",
    
    "MEHRHOLZ, J. et al. Electromechanical and robot-assisted arm training for improving activities of daily living, arm function, and arm muscle strength after stroke. Cochrane Database of Systematic Reviews, v. 2018, n. 9, 2018. Disponível em: https://www.cochranelibrary.com/cdsr/doi/10.1002/14651858.CD006876.pub5/full. Acesso em: 18 jul. 2026.",
    
    "META. Setup hand tracking in Unreal Engine. Meta Developers, 2026. Disponível em: https://developers.meta.com/horizon/documentation/unreal/unreal-hand-tracking/. Acesso em: 18 jul. 2026.",
    
    "MINISTÉRIO DA SAÚDE (BRASIL). Estratégias de atenção ao paciente com acidente vascular encefálico no SUS. Brasília: MS, 2023. Disponível em: https://www.gov.br/saude/pt-br/assuntos/saude-de-a-a-z/a/acidente-vascular-cerebral-avc. Acesso em: 18 jul. 2026.",
    
    "MINISTÉRIO DA SAÚDE (BRASIL). Rede Nacional de Atenção Especializada (RENAME). Brasília: MS, 2022. Disponível em: https://www.gov.br/saude/pt-br/assuntos/atencao-especializada. Acesso em: 18 jul. 2026.",
    
    "MOVING: A multi-modal dataset of EEG signals and virtual glove hand tracking. Sensors, v. 24, n. 16, p. 5207, 2024. Disponível em: https://www.mdpi.com/1424-8220/24/16/5207. Acesso em: 18 jul. 2026.",
    
    "PFURTSCHELLER, G.; ARANIBAR, A. Evaluation of event-related desynchronization (ERD) preceding and following voluntary self-paced movement. Electroencephalography and Clinical Neurophysiology, v. 46, n. 2, p. 138-146, 1979. Disponível em: https://pubmed.ncbi.nlm.nih.gov/46855/. Acesso em: 18 jul. 2026.",
    
    "PFURTSCHELLER, G.; LOPES DA SILVA, F. H. Event-related EEG/MEG synchronization and desynchronization: basic principles. Clinical Neurophysiology, v. 110, n. 11, p. 1842-1857, 1999. Disponível em: https://pubmed.ncbi.nlm.nih.gov/10576479/. Acesso em: 18 jul. 2026.",
    
    "PISZCZ, A.; ROJEK, I.; MIKOŁAJEWSKI, D. Impact of virtual reality on brain-computer interface performance in IoT control: review of current state of knowledge. Applied Sciences, v. 14, n. 22, p. 10541, 2024. Disponível em: https://www.mdpi.com/2076-3417/14/22/10541. Acesso em: 18 jul. 2026.",
    
    "RIYAD, M.; KHALIL, M.; ADIB, A. MI-EEGNET: a novel convolutional neural network for motor imagery classification. Journal of Neuroscience Methods, v. 353, p. 109037, 2021. Disponível em: https://www.sciencedirect.com/science/article/pii/S0165027021000538. Acesso em: 18 jul. 2026.",
    
    "SALAMI, A. et al. EEG-ITNet: an explainable inception temporal convolutional network for motor imagery classification. IEEE Access, v. 10, p. 36672-36685, 2022. Disponível em: https://ieeexplore.ieee.org/document/9747584. Acesso em: 18 jul. 2026.",
    
    "SCHIRRMEISTER, R. T. et al. Deep learning with convolutional neural networks for EEG decoding and visualization. Human Brain Mapping, v. 38, n. 11, p. 5391-5420, 2017. Disponível em: https://onlinelibrary.wiley.com/doi/10.1002/hbm.23730. Acesso em: 18 jul. 2026.",
    
    "TAYLOR, R. M. et al. VRPN: a device-independent, network-transparent VR peripheral system. In: Proceedings of the ACM Symposium on Virtual Reality Software and Technology (VRST). 2001. p. 55-61. Disponível em: https://dl.acm.org/doi/10.1145/505008.505019. Acesso em: 18 jul. 2026.",
    
    "VOURVOPOULOS, A. et al. RehabNet: a distributed architecture for motor and cognitive neuro-rehabilitation. In: IEEE 15th International Conference on E-Health Networking, Applications & Services (Healthcom). 2013. p. 454-459. Disponível em: https://ieeexplore.ieee.org/document/6720703. Acesso em: 18 jul. 2026.",
    
    "VOURVOPOULOS, A.; BERMÚDEZ I BADIA, S. Motor priming in virtual reality for motor rehabilitation. In: International Conference on Virtual Rehabilitation (ICVR). 2015. Disponível em: https://ieeexplore.ieee.org/document/7458799. Acesso em: 18 jul. 2026.",
    
    "WOLPAW, J. R. et al. Brain-computer interfaces for communication and control. Clinical Neurophysiology, v. 113, n. 6, p. 767-791, 2002. Disponível em: https://pubmed.ncbi.nlm.nih.gov/12048038/. Acesso em: 18 jul. 2026.",
    
    "WOO, S. et al. An open source-based BCI application for virtual world tour and its usability evaluation. Frontiers in Human Neuroscience, v. 15, p. 675986, 2021. Disponível em: https://www.frontiersin.org/articles/10.3389/fnhum.2021.675986/full. Acesso em: 18 jul. 2026.",
    
    "ZEILER, S. R.; KRASCHI, K. R. The interaction between training and plasticity in the poststroke brain. Current Opinion in Neurology, v. 26, n. 6, p. 609-616, 2013. Disponível em: https://pubmed.ncbi.nlm.nih.gov/24136129/. Acesso em: 18 jul. 2026."
]

for ref in referencias:
    add_para(doc, ref, indent=Cm(0))

page_break(doc)

# ============================================
# ANEXO E APÊNDICE
# ============================================
add_title(doc, "ANEXO")
add_para(doc, "")
add_para(doc, "Anexo A – Termo de Consentimento Livre e Esclarecido (TCLE)")
add_para(doc, "[O TCLE completo será incluído aqui, contemplando: identificação da pesquisa, objetivos, procedimentos, riscos e benefícios, garantias de sigilo, voluntariedade da participação, dados do pesquisador responsável e do CEP.]")

page_break(doc)

add_title(doc, "APÊNDICE")
add_para(doc, "")

add_para(doc, "Apêndice A – Glossário de Termos Técnicos")
add_para(doc, "BCI – Brain-Computer Interface (Interface Cérebro-Computador)")
add_para(doc, "CNN – Convolutional Neural Network (Rede Neural Convolucional)")
add_para(doc, "EEG – Electroencephalography (Eletroencefalografia)")
add_para(doc, "ERD/ERS – Event-Related Desynchronization/Synchronization")
add_para(doc, "LSL – Lab Streaming Layer (protocolo de streaming temporal)")
add_para(doc, "MI – Motor Imagery (Imagética Motora)")
add_para(doc, "Middleware – Camada de software intermediária entre aplicações")
add_para(doc, "VR – Virtual Reality (Realidade Virtual)")
add_para(doc, "ZeroMQ – Biblioteca de mensageria de alta performance")

add_para(doc, "")
add_para(doc, "Apêndice B – Especificação da API do Middleware")
add_para(doc, "[Documentação técnica detalhada das interfaces de programação, endpoints, formatos de mensagem e exemplos de uso.]")

add_para(doc, "")
add_para(doc, "Apêndice C – Modelo de Análise de Custo-Efetividade")
add_para(doc, "[Planilha/template de análise de decisão em saúde com cenários de implantação no SUS e saúde suplementar.]")

# ============================================
# SALVAR DOCUMENTO
# ============================================
timestamp = datetime.now().strftime("%d%m%Y_%H%M")
output_file = f"Projeto_Qualificacao_Mestrado_Middleware_BCI_VR_AVE_{timestamp}.docx"
doc.save(output_file)
print(f"Documento salvo com sucesso: {output_file}")
